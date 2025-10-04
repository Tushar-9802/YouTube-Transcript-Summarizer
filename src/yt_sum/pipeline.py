# src/yt_sum/pipeline.py
from __future__ import annotations
from typing import Dict, Any, Tuple, Optional
from pathlib import Path
import re
import torch
import gc
import logging

from docx import Document
from src.yt_sum.models.transcriber import Transcriber
from src.yt_sum.models.summarizer import Summarizer
from src.yt_sum.utils.keywords import extract_keywords, highlight_sentences, extract_critical_terms
from src.yt_sum.utils.downloader import download_audio
from src.yt_sum.utils.logging import get_logger

logger = get_logger("pipeline", level="INFO")


def aggressive_cleanup():
    """Triple-pass memory cleanup for maximum VRAM release."""
    gc.collect()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
        torch.cuda.synchronize()
        torch.cuda.ipc_collect()
    gc.collect()
    logger.info("GPU memory aggressively cleared")


def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    return text.encode("utf-8", "ignore").decode("utf-8").strip()


def export_docx(base_filename: str, transcript: str, summary: str) -> str:
    out = Path(f"{base_filename}.docx")
    doc = Document()
    doc.add_heading("Video Summary", 0)
    doc.add_heading("Summary", level=1)
    for para in (summary or "").split("\n"):
        p = _clean_text(para)
        if p:
            doc.add_paragraph(p)
    doc.add_heading("Transcript", level=1)
    for para in (transcript or "").split("\n"):
        p = _clean_text(para)
        if p:
            doc.add_paragraph(p)
    doc.save(out)
    return str(out)


def run_pipeline(
    url: str,
    workdir: Path,
    *,
    domain: str = "general",
    whisper_size: Optional[str] = None,
    prefer_accuracy: bool = True,
    summarizer_model: Optional[str] = None,
    use_8bit: Optional[bool] = None,
    refinement: bool = True,
    imrad: bool = False,
    min_len: Optional[int] = None,
    max_len: Optional[int] = None,
    chunk_tokens: Optional[int] = None,
    chunk_overlap: Optional[int] = None,
    translate_non_english: bool = True,
    compression_ratio: Optional[float] = None,
    audience: str = "expert",
    output_language: Optional[str] = None,
    progress_callback=None,
) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    """
    Production pipeline optimized for 90-minute videos on 8GB VRAM.
    
    Optimizations:
    - Aggressive memory cleanup between stages
    - Mistral 7B with 4-bit quantization
    - Forced quantization on <=16GB VRAM
    - Streaming processing for long transcripts
    """
    
    # Device detection
    device_info = {"cuda": torch.cuda.is_available()}
    vram_gb = 0
    
    if torch.cuda.is_available():
        props = torch.cuda.get_device_properties(0)
        vram_gb = round(props.total_memory / (1024 ** 3), 2)
        device_info.update(name=props.name, total_gb=vram_gb)
        
        # Auto-enable quantization on <=16GB (Mistral 7B needs ~14GB FP16)
        if use_8bit is None and vram_gb <= 16:
            use_8bit = True
            logger.info(f"Auto-enabling 4-bit quantization for {vram_gb}GB VRAM")
    else:
        device_info.update(name="CPU", total_gb=None)

    logger.info(f"Pipeline starting: {device_info}")
    
    # ALWAYS use Mistral 7B (removed Phi-3 fallback)
    if not summarizer_model:
        summarizer_model = "mistralai/Mistral-7B-Instruct-v0.2"
        logger.info(f"Using Mistral 7B with {'4-bit' if use_8bit else 'FP16'} precision")
    
    # === STAGE 1: Download ===
    logger.info("Stage 1/4: Downloading audio...")
    audio_dir = workdir / "audio"
    audio_path, meta = download_audio(url, audio_dir)
    duration_sec = meta.get("duration", 0)
    
    # === STAGE 2: Transcription ===
    logger.info(f"Stage 2/4: Transcribing ({duration_sec/60:.1f} minutes)...")
    transcriber = Transcriber(model_size=whisper_size, prefer_accuracy=prefer_accuracy)
    segments = transcriber.transcribe(
        str(audio_path),
        domain=domain,
        translate_to_english=translate_non_english,
    )
    
    transcript = " ".join(_clean_text(s["text"]) for s in segments).strip()
    if not transcript:
        raise RuntimeError("Empty transcript")
    
    # Save transcript
    txt_dir = workdir / "transcripts"
    txt_dir.mkdir(parents=True, exist_ok=True)
    txt_path = txt_dir / f"{meta.get('id','video')}_transcript.txt"
    txt_path.write_text(transcript, encoding="utf-8")
    
    logger.info(f"Transcript: {len(transcript)} chars, {len(transcript.split())} words")
    
    # CRITICAL: Clear transcriber and Whisper from VRAM
    logger.info("Clearing Whisper from GPU...")
    del transcriber
    del segments
    aggressive_cleanup()
    
    if torch.cuda.is_available():
        free_gb = torch.cuda.mem_get_info()[0] / (1024**3)
        logger.info(f"Free VRAM after transcription: {free_gb:.2f} GB")
    
    # === STAGE 3: Summarization ===
    logger.info("Stage 3/4: Loading Mistral 7B summarizer...")
    
    summarizer = Summarizer(
        domain=domain,
        summarizer_model=summarizer_model,
        use_8bit=bool(use_8bit)
    )
    
    # Apply user overrides
    if chunk_tokens:
        summarizer.chunk_tokens = chunk_tokens
    if chunk_overlap:
        summarizer.chunk_overlap = chunk_overlap
    
    logger.info(f"Summarizing: chunks={summarizer.chunk_tokens}, overlap={summarizer.chunk_overlap}")
    
    summary_text = summarizer.summarize_long(
        transcript,
        imrad=imrad,
        refinement=refinement,
        min_len=min_len,
        max_len=max_len,
        chunk_tokens=summarizer.chunk_tokens,
        chunk_overlap=summarizer.chunk_overlap,
        compression_ratio=compression_ratio or 0.2,
        audience=audience,
        output_language=output_language,
        duration_seconds=duration_sec,
    )
    
    # Clear summarizer
    logger.info("Clearing summarizer from GPU...")
    summarizer.unload()
    del summarizer
    aggressive_cleanup()
    
    # === STAGE 4: Post-processing ===
    logger.info("Stage 4/4: Extracting keywords and highlights...")
    keywords = [k for k in extract_keywords(summary_text, top_n=15, method="auto", max_ngram=3) if k]
    critical_terms = [t for t in extract_critical_terms(transcript) if t]
    highlights = [_clean_text(h) for h in highlight_sentences(transcript, keywords, top_k=8, diversity=0.35)]
    
    # Export
    out_dir = workdir / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        docx_path = export_docx(str(out_dir / f"{meta.get('id','summary')}"), transcript, summary_text)
    except Exception as e:
        logger.warning(f"DOCX export failed: {e}")
        docx_path = None
    
    # Final cleanup
    aggressive_cleanup()
    
    meta_info = {
        "video": meta,
        "device": device_info,
        "paths": {
            "audio": str(audio_path.resolve()),
            "transcript": str(txt_path.resolve()),
            "docx": str(Path(docx_path).resolve()) if docx_path else None,
        },
    }
    
    results = {
        "transcript": transcript,
        "summary": summary_text,
        "keywords": keywords,
        "critical_terms": critical_terms,
        "highlights": highlights,
    }
    
    logger.info("Pipeline complete!")
    return meta_info, results