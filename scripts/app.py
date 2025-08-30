# scripts/app.py
# --- make ./src importable no matter how Streamlit starts the process ---
from pathlib import Path
import sys
ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

import streamlit as st
import torch
from docx import Document

from yt_sum.utils.config import Config
from yt_sum.pipeline import run_pipeline

st.set_page_config(page_title="YouTube Transcript Summarizer", layout="wide")
st.title("YouTube Transcript Summarizer")

cfg = Config("configs/default.yaml")
cfg.ensure_dirs()

# -----------------------
# Sidebar: Inputs/Controls
# -----------------------
with st.sidebar:
    st.header("Settings")

    youtube_url = st.text_input(
        "YouTube URL",
        placeholder="https://www.youtube.com/watch?v=...",
        help="Paste the full video URL. The app will fetch metadata, audio, and transcript (via Whisper if needed)."
    )

    # Preset models
    summarizer_model = st.selectbox(
        "Summarizer Model",
        ["facebook/bart-large-cnn", "t5-large"],  # keep BART as default
        index=0,
        help="Single strong general model. BART CNN/DM is a robust baseline."
    )

    custom_model = st.text_input(
        "Custom HF model ID (optional)",
        placeholder="e.g., allenai/led-large-16384",
        help="Override the preset with any Hugging Face seq2seq summarization model (supports AutoModelForSeq2SeqLM)."
    )
    chosen_model = (custom_model.strip() or summarizer_model)

    # Core summarization controls
    st.subheader("Summary Controls")
    final_tokens = st.slider(
        "Final summary size (tokens)",
        min_value=160, max_value=600, value=320, step=20,
        help="Target length for the FINAL fused summary.\n"
             "Tip: 280–360 is a good paper-style abstract; raise for long lectures."
    )
    per_chunk_ratio = st.slider(
        "Per-chunk ratio",
        min_value=0.08, max_value=0.35, value=0.18, step=0.01,
        help="Approximate fraction of each chunk to keep in the first-pass summaries. "
             "Higher = more detail carried into the final fuse."
    )
    beams = st.slider(
        "Detail level (num_beams)",
        min_value=1, max_value=8, value=5, step=1,
        help="Beam search width. Higher = more thorough but slower. 4–6 is a solid default."
    )
    chunk_overlap = st.slider(
        "Chunk overlap (tokens)",
        min_value=64, max_value=256, value=120, step=16,
        help="How much the chunks overlap to avoid cutting sentences. "
             "Increase slightly if you see context being dropped between chunks."
    )
    chunked = st.checkbox(
        "Chunked summarization (recommended)",
        value=True,
        help="Enable map→reduce summarization for long transcripts. Prevents truncation and improves coverage."
    )

    # Decoding style
    st.subheader("Decoding Style")
    use_sampling = st.checkbox(
        "Enable sampling (more creative)",
        value=False,
        help="Turns on sampling for more varied phrasing. For research/reproducibility, leave OFF."
    )
    temperature = st.slider(
        "Temperature",
        min_value=0.2, max_value=1.5, value=0.9, step=0.1,
        help="Higher = more randomness. Ignored if sampling is OFF.",
        disabled=not use_sampling
    )

    # Efficiency / precision
    st.subheader("Performance & ASR")
    use_8bit = st.checkbox(
        "Use 8-bit (GPU) quantization",
        value=True,
        help="Loads the summarizer in 8-bit to save VRAM with minimal quality loss. "
             "Requires a CUDA GPU. Falls back automatically if unavailable."
    )
    prefer_accuracy = st.checkbox(
        "Prefer accuracy over speed (auto ASR)",
        value=True,
        help="Lets the pipeline choose a larger Whisper model / safer compute for better transcription quality."
    )
    force_whisper_size = st.selectbox(
        "ASR size override (optional)",
        [None, "tiny", "base", "small", "medium", "large-v3"],
        index=0,
        help="Force a specific Whisper size. Leave at None for the app to auto-pick based on GPU/length."
    )
    language = st.selectbox(
        "Force language (optional)",
        [None, "en", "hi", "fr", "de", "es"],
        index=0,
        help="If the video language is known, set it here to speed up and improve ASR. Otherwise leave None."
    )

    # Guided / domain-aware hook (already usable)
    st.subheader("Guided Summary (optional)")
    guidance = st.text_area(
        "Guidance / focus",
        placeholder="e.g., Produce a sectioned summary: Background, Methods, Results, Limitations.",
        height=90,
        help="Add light instructions to steer the summary. Works best with BART/LED/T5. "
             "Keep it concise and factual for research outputs."
    )

# In-UI explanations / docs
with st.expander("What 'Auto' does & control guide"):
    st.markdown(
        """
**Auto choices:**  
- Detects GPU & free VRAM, then chooses a suitable **Faster-Whisper** size and compute type.  
- Tunes **chunk size/overlap**, **beam width**, and **length budgets** relative to transcript length, aiming for stable coverage.

**Model guidance:**  
- **BART CNN/DM** → strong general baseline, rich phrasing.  
- **LED-base-16384** → handles very long inputs (up to 16k tokens); works well with map→reduce too.  
- **LED-large-16384-arxiv** → biased toward scientific writing and arXiv-style summaries.  
- **SCITLDR (citation-aware)** → concise scientific TL;DRs; increase target length if you need more detail.  
- **T5-large** → flexible; good with guided/sectioned prompts.

**Controls:**  
- **Final summary size (tokens):** Target length for the fused output.  
- **Per-chunk ratio:** Portion of each chunk kept during the first pass (higher → more detail).  
- **Detail level (num_beams):** Wider beam search improves thoroughness but costs time.  
- **Chunk overlap:** Overlaps chunk boundaries to avoid context loss.  
- **Sampling & temperature:** For creative phrasing; keep off for reproducible research.  
- **8-bit (GPU):** Saves VRAM with little quality loss on modern GPUs.  
- **Guidance:** Add a short instruction to shape structure/emphasis (e.g., Methods/Results/Limitations).
        """
    )

# ---------------
# Progress display
# ---------------
stage = st.empty()
progress = st.progress(0)

def progress_cb(which: str, pct: int, extra):
    names = {"init": "Init", "download": "Downloading", "transcribe": "Transcribing", "summarize": "Summarizing"}
    stage.info(f"{names.get(which, which)}… {pct}%")
    progress.progress(pct)
    if which == "init":
        dev = extra.get("device", {})
        st.caption(f"Device: {dev.get('name')} (CUDA={dev.get('cuda')}) • Free {dev.get('free_gb')} / {dev.get('total_gb')} GB")

# --------------
# DOCX export fn
# --------------
def export_docx(title: str, transcript: str, summary: str) -> Path:
    out = cfg.outputs_dir / f"{title}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("YouTube Transcript Summary", level=0)
    doc.add_paragraph(f"Title: {title}")
    doc.add_paragraph("")
    doc.add_heading("Summary", level=1); doc.add_paragraph(summary)
    doc.add_heading("Transcript", level=1); doc.add_paragraph(transcript)
    doc.save(out)
    return out
# -------------------------
# Storage & Cleanup (new)
# -------------------------
import os
import time
from datetime import datetime

def _human_bytes(n: int) -> str:
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if n < 1024:
            return f"{n:.1f} {unit}"
        n /= 1024.0
    return f"{n:.1f} PB"

def _dir_size(p: Path) -> int:
    total = 0
    if not p.exists():
        return 0
    for root, _, files in os.walk(p):
        for fn in files:
            fp = Path(root) / fn
            try:
                total += fp.stat().st_size
            except Exception:
                pass
    return total

def _list_files(p: Path, exts: tuple[str, ...] | None = None) -> list[Path]:
    if not p.exists():
        return []
    files = []
    for root, _, fs in os.walk(p):
        for fn in fs:
            fp = Path(root) / fn
            if exts is None or fp.suffix.lower() in exts:
                files.append(fp)
    # newest first
    files.sort(key=lambda q: q.stat().st_mtime if q.exists() else 0, reverse=True)
    return files

def _label_for(fp: Path) -> str:
    try:
        sz = _human_bytes(fp.stat().st_size)
        ts = datetime.fromtimestamp(fp.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
        return f"{fp.name}  •  {sz}  •  modified {ts}"
    except Exception:
        return fp.name

st.markdown("---")
with st.expander("Storage & Cleanup"):
    # Sizes
    audio_size = _dir_size(cfg.audio_dir)
    trans_size = _dir_size(cfg.transcript_dir)
    out_size = _dir_size(cfg.outputs_dir)

    st.write({
        "Audio directory": {"path": str(cfg.audio_dir), "size": _human_bytes(audio_size)},
        "Transcripts directory": {"path": str(cfg.transcript_dir), "size": _human_bytes(trans_size)},
        "Exports directory": {"path": str(cfg.outputs_dir), "size": _human_bytes(out_size)},
    })

    tabs = st.tabs(["Audio files", "Transcripts", "Exports"])

    # --- Audio tab ---
    with tabs[0]:
        audio_exts = (".wav", ".m4a", ".mp3", ".webm", ".opus", ".mp4", ".aac", ".ogg", ".flac", ".mka", ".mkv")
        audio_files = _list_files(cfg.audio_dir, exts=audio_exts)
        if not audio_files:
            st.info("No audio files found.")
        else:
            labels = [_label_for(p) for p in audio_files]
            label_to_path = dict(zip(labels, audio_files))
            sel = st.multiselect("Select audio files to delete", options=labels, default=[])
            confirm = st.checkbox("I understand this will permanently delete the selected audio files.", key="confirm_audio")
            colA, colB = st.columns([1,1])
            with colA:
                if st.button("Delete selected audio"):
                    if not sel:
                        st.warning("No files selected.")
                    elif not confirm:
                        st.warning("Please confirm deletion.")
                    else:
                        before = _dir_size(cfg.audio_dir)
                        deleted, freed = 0, 0
                        for lbl in sel:
                            fp = label_to_path.get(lbl)
                            if not fp: 
                                continue
                            try:
                                sz = fp.stat().st_size
                                fp.unlink()
                                deleted += 1
                                freed += sz
                            except Exception as e:
                                st.error(f"Failed to delete {fp.name}: {e}")
                        after = _dir_size(cfg.audio_dir)
                        st.success(f"Deleted {deleted} file(s). Freed ~{_human_bytes(freed)}. Current audio dir size: {_human_bytes(after)}")
                        st.experimental_rerun()
            with colB:
                if st.button("Delete all audio files"):
                    if not confirm:
                        st.warning("Please confirm deletion.")
                    else:
                        before = _dir_size(cfg.audio_dir)
                        deleted, freed = 0, 0
                        for fp in audio_files:
                            try:
                                sz = fp.stat().st_size
                                fp.unlink()
                                deleted += 1
                                freed += sz
                            except Exception as e:
                                st.error(f"Failed to delete {fp.name}: {e}")
                        after = _dir_size(cfg.audio_dir)
                        st.success(f"Deleted {deleted} file(s). Freed ~{_human_bytes(freed)}. Current audio dir size: {_human_bytes(after)}")
                        st.experimental_rerun()

    # --- Transcripts tab ---
    with tabs[1]:
        txt_files = _list_files(cfg.transcript_dir, exts=(".txt",))
        if not txt_files:
            st.info("No transcript files found.")
        else:
            labels = [_label_for(p) for p in txt_files]
            label_to_path = dict(zip(labels, txt_files))
            sel = st.multiselect("Select transcripts to delete", options=labels, default=[], key="sel_txt")
            confirm = st.checkbox("I understand this will permanently delete the selected transcripts.", key="confirm_txt")
            colA, colB = st.columns([1,1])
            with colA:
                if st.button("Delete selected transcripts"):
                    if not sel:
                        st.warning("No files selected.")
                    elif not confirm:
                        st.warning("Please confirm deletion.")
                    else:
                        before = _dir_size(cfg.transcript_dir)
                        deleted, freed = 0, 0
                        for lbl in sel:
                            fp = label_to_path.get(lbl)
                            if not fp:
                                continue
                            try:
                                sz = fp.stat().st_size
                                fp.unlink()
                                deleted += 1
                                freed += sz
                            except Exception as e:
                                st.error(f"Failed to delete {fp.name}: {e}")
                        after = _dir_size(cfg.transcript_dir)
                        st.success(f"Deleted {deleted} file(s). Freed ~{_human_bytes(freed)}. Current transcripts dir size: {_human_bytes(after)}")
                        st.experimental_rerun()
            with colB:
                if st.button("Delete all transcripts"):
                    if not confirm:
                        st.warning("Please confirm deletion.")
                    else:
                        before = _dir_size(cfg.transcript_dir)
                        deleted, freed = 0, 0
                        for fp in txt_files:
                            try:
                                sz = fp.stat().st_size
                                fp.unlink()
                                deleted += 1
                                freed += sz
                            except Exception as e:
                                st.error(f"Failed to delete {fp.name}: {e}")
                        after = _dir_size(cfg.transcript_dir)
                        st.success(f"Deleted {deleted} file(s). Freed ~{_human_bytes(freed)}. Current transcripts dir size: {_human_bytes(after)}")
                        st.experimental_rerun()

    # --- Exports tab ---
    with tabs[2]:
        docx_files = _list_files(cfg.outputs_dir, exts=(".docx",))
        if not docx_files:
            st.info("No export files found.")
        else:
            labels = [_label_for(p) for p in docx_files]
            label_to_path = dict(zip(labels, docx_files))
            sel = st.multiselect("Select exports to delete", options=labels, default=[], key="sel_docx")
            confirm = st.checkbox("I understand this will permanently delete the selected exports.", key="confirm_docx")
            colA, colB = st.columns([1,1])
            with colA:
                if st.button("Delete selected exports"):
                    if not sel:
                        st.warning("No files selected.")
                    elif not confirm:
                        st.warning("Please confirm deletion.")
                    else:
                        before = _dir_size(cfg.outputs_dir)
                        deleted, freed = 0, 0
                        for lbl in sel:
                            fp = label_to_path.get(lbl)
                            if not fp:
                                continue
                            try:
                                sz = fp.stat().st_size
                                fp.unlink()
                                deleted += 1
                                freed += sz
                            except Exception as e:
                                st.error(f"Failed to delete {fp.name}: {e}")
                        after = _dir_size(cfg.outputs_dir)
                        st.success(f"Deleted {deleted} file(s). Freed ~{_human_bytes(freed)}. Current exports dir size: {_human_bytes(after)}")
                        st.experimental_rerun()
            with colB:
                if st.button("Delete all exports"):
                    if not confirm:
                        st.warning("Please confirm deletion.")
                    else:
                        before = _dir_size(cfg.outputs_dir)
                        deleted, freed = 0, 0
                        for fp in docx_files:
                            try:
                                sz = fp.stat().st_size
                                fp.unlink()
                                deleted += 1
                                freed += sz
                            except Exception as e:
                                st.error(f"Failed to delete {fp.name}: {e}")
                        after = _dir_size(cfg.outputs_dir)
                        st.success(f"Deleted {deleted} file(s). Freed ~{_human_bytes(freed)}. Current exports dir size: {_human_bytes(after)}")
                        st.experimental_rerun()


# -----
# RUN
# -----
if st.button("Run"):
    if not youtube_url.strip():
        st.error("Enter a YouTube URL")
        st.stop()

    st.info(f"Running on: {'CUDA' if torch.cuda.is_available() else 'CPU'} • Model: {chosen_model}")

    # Package summarizer (map→reduce) parameters for the pipeline / summarizer
    summarizer_params = {
        # length control
        "reduce_target_tokens": int(final_tokens),
        "summary_ratio": float(per_chunk_ratio),
        # decoding
        "num_beams": int(beams),
        "no_repeat_ngram_size": 3,
        "length_penalty": 0.9,
        "repetition_penalty": 1.08,
        "do_sample": bool(use_sampling),
        "temperature": float(temperature) if use_sampling else 1.0,
        # chunking
        "chunk_tokens": 900,          # keep conservative; LED can handle more if you bump this later
        "chunk_overlap": int(chunk_overlap),
        # guided
        "guidance": guidance or None,
        # efficiency
        "use_8bit": bool(use_8bit),
    }

    meta, res = run_pipeline(
        url=youtube_url,
        cfg=cfg,
        whisper_size=force_whisper_size,   # None => auto
        summarizer_model=chosen_model,
        chunked=chunked,
        language=language,
        prefer_accuracy=prefer_accuracy,
        summarizer_params=summarizer_params,
        progress=progress_cb,
    )

    stage.success("Done."); progress.progress(100)
    title = meta["video"]["title"] or meta["video"]["id"]

    with st.expander("Auto Choices"):
        sel = meta.get("selections", {})
        st.write({
            "Device": sel.get("device", {}),
            "ASR": sel.get("asr_choice", {}),
            "Summarizer": {
                **sel.get("summarizer", {}),
                "final_target_tokens": final_tokens,
                "per_chunk_ratio": per_chunk_ratio,
                "num_beams": beams,
                "chunk_overlap": chunk_overlap,
                "use_8bit": use_8bit,
                "sampling": use_sampling,
                "temperature": temperature if use_sampling else None,
                "guidance": (guidance[:80] + "…") if guidance and len(guidance) > 80 else guidance,
                "model": chosen_model,
            },
        })

    st.subheader("Summary")
    st.text_area("Summary Output", res["summary"], height=260)

    st.subheader("Transcript")
    st.text_area("Full Transcript", res["transcript"], height=260)

    if st.button("Export as DOCX"):
        safe = (title or "video").replace(":", " ").replace("/", " ")
        out = export_docx(safe, res["transcript"], res["summary"])
        st.success(f"Saved: {out}")
        with open(out, "rb") as f:
            st.download_button(
                "Download DOCX",
                f,
                file_name=Path(out).name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
